import os
import requests
from docx import Document
import re
from bs4 import BeautifulSoup
import tempfile
import io

# 配置分割长度变量
MAX_FRAGMENT_LENGTH = 10000

def download_word_from_url(url):
    """从URL下载Word文件"""
    try:
        response = requests.get(url)
        response.raise_for_status()
        return response.content
    except Exception as e:
        raise Exception(f"下载Word文件失败: {e}")

def word_to_html_with_styles(doc):
    """将Word文档转换为HTML，保留所有样式（不包含HTML头部和body标签）"""
    html_content = []
    
    # 获取文档中的所有元素（段落和表格）并保持原有顺序
    document_elements = get_document_elements_in_order(doc)
    
    # 按顺序处理每个元素
    for element in document_elements:
        if element['type'] == 'paragraph':
            paragraph = element['content']
            if paragraph.text.strip():
                # 获取段落样式
                style_attrs = get_paragraph_style(paragraph)
                
                # 处理段落中的文本样式
                html_paragraph = '<p'
                for attr, value in style_attrs.items():
                    html_paragraph += f' {attr}="{value}"'
                html_paragraph += '>'
                
                # 处理段落中的run样式
                for run in paragraph.runs:
                    if run.text.strip():
                        run_styles = get_run_style(run)
                        if run_styles:
                            # 如果有样式，添加span标签
                            style_str = ''
                            for attr, value in run_styles.items():
                                style_str += f'{attr}: {value}; '
                            html_paragraph += f'<span style="{style_str.strip()}">{escape_html(run.text)}</span>'
                        else:
                            html_paragraph += escape_html(run.text)
                
                html_paragraph += '</p>'
                html_content.append(html_paragraph)
        
        elif element['type'] == 'table':
            table = element['content']
            html_table = '<table border="1" style="border-collapse: collapse;">'
            for row in table.rows:
                html_table += '<tr>'
                for cell in row.cells:
                    html_table += '<td style="border: 1px solid #ddd; padding: 8px;">'
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            # 处理表格单元格中的文本样式
                            cell_text = ''
                            for run in paragraph.runs:
                                if run.text.strip():
                                    run_styles = get_run_style(run)
                                    if run_styles:
                                        style_str = ''
                                        for attr, value in run_styles.items():
                                            style_str += f'{attr}: {value}; '
                                        cell_text += f'<span style="{style_str.strip()}">{escape_html(run.text)}</span>'
                                    else:
                                        cell_text += escape_html(run.text)
                            html_table += cell_text
                    html_table += '</td>'
                html_table += '</tr>'
            html_table += '</table>'
            html_content.append(html_table)
    
    return '\n'.join(html_content)

def get_document_elements_in_order(doc):
    """获取文档中的所有元素（段落和表格）并保持原有顺序"""
    elements = []
    
    # 获取所有段落的索引位置
    paragraph_indices = {}
    for i, paragraph in enumerate(doc.paragraphs):
        paragraph_indices[paragraph._element] = i
    
    # 获取所有表格的索引位置
    table_indices = {}
    for i, table in enumerate(doc.tables):
        table_indices[table._element] = i
    
    # 遍历文档的所有元素，按照它们在XML中的顺序
    for element in doc.element.body:
        if element in paragraph_indices:
            # 这是一个段落
            paragraph_index = paragraph_indices[element]
            elements.append({
                'type': 'paragraph',
                'content': doc.paragraphs[paragraph_index]
            })
        elif element in table_indices:
            # 这是一个表格
            table_index = table_indices[element]
            elements.append({
                'type': 'table',
                'content': doc.tables[table_index]
            })
    
    return elements

def get_paragraph_style(paragraph):
    """获取段落的样式属性"""
    styles = {}
    
    if paragraph.style.name:
        styles['class'] = paragraph.style.name.lower().replace(' ', '-')
    
    # 获取对齐方式
    alignment = paragraph.alignment
    if alignment:
        alignment_map = {
            1: 'left',
            2: 'center',
            3: 'right',
            4: 'justify'
        }
        if alignment in alignment_map:
            styles['align'] = alignment_map[alignment]
    
    return styles

def get_run_style(run):
    """获取run的样式属性"""
    styles = {}
    
    if run.bold:
        styles['font-weight'] = 'bold'
    if run.italic:
        styles['font-style'] = 'italic'
    if run.underline:
        styles['text-decoration'] = 'underline'
    
    # 获取字体
    if run.font.name:
        styles['font-family'] = run.font.name
    
    # 获取字号
    if run.font.size:
        # Word中字号是以half-points为单位，需要转换为pt
        font_size_pt = run.font.size / 2
        styles['font-size'] = f'{font_size_pt}pt'
    
    # 获取颜色
    if run.font.color and run.font.color.rgb:
        rgb = run.font.color.rgb
        styles['color'] = f'#{rgb:06x}'
    
    return styles

def escape_html(text):
    """转义HTML特殊字符"""
    return (text.replace('&', '&amp;')
               .replace('<', '&lt;')
               .replace('>', '&gt;')
               .replace('"', '&quot;')
               .replace("'", '&#39;'))

def is_heading_tag(html_content, position):
    """检查指定位置是否是标题标签的结尾"""
    # 查找position附近的标签
    start_pos = max(0, position - 100)
    end_pos = min(len(html_content), position + 100)
    context = html_content[start_pos:end_pos]
    
    # 检查是否有h1-h6标签
    heading_pattern = r'<h[1-6][^>]*>.*?</h[1-6]>'
    matches = re.finditer(heading_pattern, context, re.DOTALL)
    
    for match in matches:
        match_start = start_pos + match.start()
        match_end = start_pos + match.end()
        if match_start <= position <= match_end:
            return True
    
    return False

def find_safe_split_point(html_content, max_length):
    """找到安全的分割点，确保不在标题处分割且不在标签中间切割"""
    if len(html_content) <= max_length:
        return len(html_content)
    
    # 从max_length位置开始向前查找合适的分割点
    split_point = max_length
    
    # 避免在标题处分割
    while split_point > max_length - 500 and split_point > 0:
        if not is_heading_tag(html_content, split_point):
            # 查找段落结束标签
            paragraph_end = html_content.rfind('</p>', 0, split_point)
            if paragraph_end != -1 and paragraph_end > max_length - 500:
                split_point = paragraph_end + 4  # 包含</p>标签
                break
        split_point -= 1
    
    # 如果没找到合适的段落结束，查找其他标签结束
    if split_point == max_length:
        for i in range(max_length, max(max_length - 500, 0), -1):
            if html_content[i] == '>' and not is_heading_tag(html_content, i):
                split_point = i + 1
                break
    
    # 确保不在标签中间切割
    split_point = ensure_not_in_tag_middle(html_content, split_point)
    
    # 检查分割点是否在标题标签开始处，如果是则提前分割点
    split_point = avoid_heading_tag_at_split_point(html_content, split_point)
    
    return split_point

def avoid_heading_tag_at_split_point(html_content, split_point):
    """避免标题标签出现在分割点处，将整个标题标签移到下个片段"""
    if split_point <= 0 or split_point >= len(html_content):
        return split_point
    
    # 检查分割点附近是否有标题标签的开始
    # 向前查找最近的标签开始
    for i in range(split_point - 1, max(-1, split_point - 200), -1):
        if i < 0:
            break
        if html_content[i] == '<':
            # 检查是否是标题相关的标签（如 p class="heading-1"）
            tag_content = html_content[i:split_point]
            if is_heading_tag_start(tag_content):
                # 找到标题标签的开始，需要找到对应的结束标签
                tag_end = html_content.find('>', i)
                if tag_end != -1:
                    # 找到标题标签的结束位置
                    heading_end_tag = html_content.find('</p>', tag_end)
                    if heading_end_tag != -1:
                        # 如果整个标题标签在分割点附近，将分割点移到标题标签开始前
                        if i < split_point <= heading_end_tag + 4:
                            return i
                break
    
    return split_point

def is_heading_tag_start(tag_content):
    """检查标签内容是否是标题标签的开始"""
    # 检查是否包含 heading-1, heading-2, heading-3 等类名
    heading_patterns = ['heading-1', 'heading-2', 'heading-3', 'heading-4', 'heading-5', 'heading-6']
    for pattern in heading_patterns:
        if pattern in tag_content:
            return True
    return False

def ensure_not_in_tag_middle(html_content, split_point):
    """确保分割点不在标签中间，只在标签开始前或结束后切割"""
    if split_point <= 0 or split_point >= len(html_content):
        return split_point
    
    # 检查分割点是否在标签中间
    # 向前查找最近的标签开始或结束
    for i in range(split_point - 1, max(-1, split_point - 100), -1):
        if i < 0:
            break
        if html_content[i] == '<':
            # 找到了标签开始，检查是否是结束标签
            if i + 1 < len(html_content) and html_content[i + 1] == '/':
                # 这是结束标签的开始，可以在这个位置之前切割
                return i
            else:
                # 这是开始标签的开始，需要找到对应的结束标签
                tag_end = html_content.find('>', i)
                if tag_end != -1:
                    # 如果分割点在标签中间，移动到标签结束后
                    if i < split_point <= tag_end:
                        return tag_end + 1
                break
        elif html_content[i] == '>':
            # 找到了标签结束，可以在这个位置之后切割
            return i + 1
    
    # 向后查找最近的标签开始或结束
    for i in range(split_point, min(len(html_content), split_point + 100)):
        if html_content[i] == '<':
            # 找到了标签开始，可以在这个位置之前切割
            return i
        elif html_content[i] == '>':
            # 找到了标签结束，可以在这个位置之后切割
            return i + 1
    
    return split_point

def split_html_content(html_content, max_length=MAX_FRAGMENT_LENGTH):
    """将HTML内容分割成指定长度的片段"""
    fragments = []
    remaining_content = html_content
    
    while remaining_content:
        if len(remaining_content) <= max_length:
            fragments.append(remaining_content)
            break
        
        # 找到安全的分割点
        split_point = find_safe_split_point(remaining_content, max_length)
        
        # 确保分割点不会太短
        if split_point < max_length * 0.5:
            split_point = max_length
        
        fragment = remaining_content[:split_point]
        fragments.append(fragment)
        remaining_content = remaining_content[split_point:]
        
        print(f"分割片段: 长度 {len(fragment)}, 剩余长度: {len(remaining_content)}")
    
    return fragments

def word_to_html_array(url, max_length=MAX_FRAGMENT_LENGTH):
    """主函数：将Word文档从URL转换为HTML数组"""
    print(f"开始处理URL: {url}")
    print(f"最大片段长度: {max_length}")
    
    # 下载Word文件
    print("正在下载Word文件...")
    word_content = download_word_from_url(url)
    
    # 解析Word文档
    print("正在解析Word文档...")
    # 使用临时文件来处理字节数据
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        temp_file.write(word_content)
        temp_file_path = temp_file.name
    
    try:
        doc = Document(temp_file_path)
    finally:
        # 清理临时文件
        os.unlink(temp_file_path)
    
    # 转换为HTML
    print("正在转换为HTML...")
    html_content = word_to_html_with_styles(doc)
    print(f"HTML总长度: {len(html_content)}")
    
    # 分割HTML内容
    print("正在分割HTML内容...")
    html_fragments = split_html_content(html_content, max_length)
    
    print(f"分割完成，共生成{len(html_fragments)}个片段")
    
    return html_fragments